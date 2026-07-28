import io
import logging

from fastapi import UploadFile
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SentenceSplitter

from app.core.clients import ocr_client, openrouter_client, s3_client
from app.core.config import settings
from app.core.exceptions import AppException, BadRequestError
from app.core.logger import IngestionLogger
from app.repositories.document import DocumentRepository
from app.schemas.document import DocumentResponse

logger = logging.getLogger(__name__)


class IngestionService:
    """Service handling document upload, S3 storage, OCR extraction, LlamaIndex chunking, embedding, and storage."""

    def __init__(self, doc_repo: DocumentRepository):
        self.doc_repo = doc_repo

    async def upload_and_process(
        self,
        file: UploadFile,
        tenant_id: str | None = None,
        user_id: str | None = None,
        folder_id: str | None = None,
        folder_path: str = "/",
    ) -> DocumentResponse:
        """Pipeline for processing document uploads."""
        filename = file.filename or "unnamed_file"
        file_ext = filename.split(".")[-1].lower() if "." in filename else ""

        if file_ext not in {"pdf", "png", "jpg", "jpeg", "docx", "xlsx"}:
            raise BadRequestError(
                f"Unsupported file type '{file_ext}'. Allowed types: pdf, png, jpg, jpeg, docx, xlsx"
            )

        if folder_id in ("root", "null", "", "undefined"):
            folder_id = None

        # Duplicate file name validation in folder
        from sqlalchemy import select

        from app.models.document import Document

        if folder_id is None:
            stmt = select(Document).where(
                Document.name == filename,
                Document.tenant_id == tenant_id,
                Document.folder_id.is_(None),
            )
        else:
            stmt = select(Document).where(
                Document.name == filename,
                Document.tenant_id == tenant_id,
                Document.folder_id == folder_id,
            )

        existing_doc = await self.doc_repo.db.execute(stmt)
        if existing_doc.scalar_one_or_none():
            folder_info = f"folder '{folder_path}'" if folder_id else "this folder"
            raise BadRequestError(
                f"A document named '{filename}' already exists in {folder_info}. Please rename or delete the existing file."
            )

        file_bytes = await file.read()
        file_size = len(file_bytes)
        print(f"[INGESTION STEP 1] Read {file_size} bytes for '{filename}'")

        # 1. Upload original file to S3
        IngestionLogger.step_header(1, "Uploading File to S3 Storage")
        s3_key = f"documents/{filename}"
        s3_url = await s3_client.upload_file(file_bytes, s3_key)
        IngestionLogger.log_s3_upload(filename, file_size, s3_key, s3_url)
        print(f"[INGESTION S3 OK] s3_url={s3_url}")

        # 2. Create initial Document entity in database
        doc = await self.doc_repo.create(
            {
                "name": filename,
                "file_type": file_ext,
                "s3_key": s3_key,
                "s3_url": s3_url,
                "file_size": file_size,
                "status": "processing",
                "tenant_id": tenant_id,
                "folder_id": folder_id,
                "folder_path": folder_path,
                "error_message": None,
            }
        )
        print(f"[INGESTION DB DOC CREATED] doc_id={doc.id}, name={doc.name}")

        # 3. Process document text, chunking via LlamaIndex, embedding, and storage
        try:
            IngestionLogger.step_header(2, "Extracting Text & Document Annotations via OCR")
            extracted_chunks, doc_metadata = await self._extract_and_chunk(
                file_bytes, filename, file_ext
            )
            print(f"[INGESTION OCR/CHUNKING OK] extracted {len(extracted_chunks)} text chunks")

            if not extracted_chunks:
                raise AppException(
                    message="No text content could be extracted from the document",
                    code="EMPTY_DOCUMENT",
                )

            # Generate embeddings
            IngestionLogger.step_header(4, "Generating OpenRouter Vector Embeddings")
            texts = [c[0] for c in extracted_chunks]
            embeddings = await openrouter_client.get_embeddings(texts)
            print(f"[INGESTION EMBEDDINGS OK] generated {len(embeddings)} vector embeddings")

            # Build chunk entities with rich metadata
            IngestionLogger.step_header(5, "Saving Document Chunks to PostgreSQL Database")
            chunk_records = []
            for idx, ((content, page_num), embedding) in enumerate(
                zip(extracted_chunks, embeddings)
            ):
                meta_dict = {
                    "doc_name": filename,
                    "page": page_num,
                    "doc_type": doc_metadata.get("document_type", "document"),
                    "short_description": doc_metadata.get("short_description", ""),
                    "summary": doc_metadata.get("summary", ""),
                    "chunk_index": idx,
                }
                chunk_records.append(
                    {
                        "document_id": doc.id,
                        "content": content,
                        "chunk_index": idx,
                        "page_number": page_num,
                        "metadata": meta_dict,
                        "embedding": embedding,
                    }
                )

            await self.doc_repo.create_chunks(chunk_records)
            doc = await self.doc_repo.update(doc, {"status": "indexed", "error_message": None})
            IngestionLogger.log_db_saved(doc.id, len(chunk_records))
            print(f"[INGESTION COMPLETED SUCCESS] doc_id={doc.id} set to INDEXED")

            # 6. Log Document Ingestion Usage & Cost (OCR + Embeddings)
            if tenant_id:
                try:
                    from app.models.query_log import QueryLog

                    prompt_t = 500 * max(1, len(chunk_records))
                    comp_t = 100 * max(1, len(chunk_records))
                    cost = round(0.0012 * max(1, len(chunk_records)), 6)
                    ingest_log = QueryLog(
                        tenant_id=tenant_id,
                        user_id=user_id,
                        question=f"Document Ingestion: {filename}",
                        answer=f"Indexed {len(chunk_records)} chunks via Mistral OCR & OpenRouter embeddings",
                        model_name="mistral-ocr+openrouter-embed",
                        prompt_tokens=prompt_t,
                        completion_tokens=comp_t,
                        total_tokens=prompt_t + comp_t,
                        estimated_cost=cost,
                        latency_ms=1200,
                        top_k=0,
                        sources_count=len(chunk_records),
                    )
                    self.doc_repo.db.add(ingest_log)
                    await self.doc_repo.db.flush()
                except Exception as e:
                    logger.warning(f"Failed to log ingestion cost: {e}")
                    print(f"[INGESTION COST LOG WARNING] {e}")

        except Exception as e:
            logger.error(f"Ingestion processing failed for document {doc.id}: {e}")
            err_msg = str(e)
            print(f"[INGESTION PROCESSING ERROR] doc_id={doc.id}, error={err_msg}")
            doc = await self.doc_repo.update(doc, {"status": "failed", "error_message": err_msg})

        return DocumentResponse.model_validate(doc)

    async def _extract_and_chunk(
        self, file_bytes: bytes, filename: str, file_ext: str
    ) -> tuple[list[tuple[str, int]], dict]:
        """Extracts text by file type, processes nodes via LlamaIndex SentenceSplitter, and returns metadata."""
        pages_content: list[tuple[str, int]] = []
        doc_annotation: dict = {}

        if file_ext in {"png", "jpg", "jpeg"}:
            text, doc_annotation = await ocr_client.extract_text_and_metadata(file_bytes, filename)
            pages_content.append((text, 1))

        elif file_ext == "pdf":
            # 1. Try Mistral OCR first if API key is configured
            if settings.MISTRAL_OCR_API_KEY and not settings.MISTRAL_OCR_API_KEY.startswith(
                "your-"
            ):
                try:
                    text, doc_annotation = await ocr_client.extract_text_and_metadata(
                        file_bytes, filename
                    )
                    if text and text.strip():
                        pages_content.append((text.strip(), 1))
                except Exception as ocr_err:
                    logger.warning(
                        f"Mistral OCR failed for {filename}, falling back to native pypdf: {ocr_err}"
                    )

            # 2. Native PDF fallback if OCR was skipped or returned empty
            if not pages_content:
                try:
                    import pypdf

                    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                    for page_idx, page in enumerate(reader.pages, 1):
                        txt = page.extract_text() or ""
                        if txt.strip():
                            pages_content.append((txt.strip(), page_idx))
                except Exception as e:
                    logger.warning(f"Native PDF parsing failed for {filename}: {e}")

        elif file_ext == "docx":
            import docx

            doc_file = docx.Document(io.BytesIO(file_bytes))
            full_text = "\n".join([p.text for p in doc_file.paragraphs if p.text.strip()])
            pages_content.append((full_text, 1))

        elif file_ext == "xlsx":
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            sheet_texts = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    row_str = " | ".join([str(cell) for cell in row if cell is not None])
                    if row_str.strip():
                        rows.append(row_str)
                if rows:
                    sheet_texts.append(f"--- Sheet: {sheet} ---\n" + "\n".join(rows))
            pages_content.append(("\n\n".join(sheet_texts), 1))

        # Perform chunking via LlamaIndex SentenceSplitter node parser
        IngestionLogger.step_header(
            3, "Chunking via LlamaIndex SentenceSplitter (Size: 1000, Overlap: 150)"
        )
        splitter = SentenceSplitter(chunk_size=1000, chunk_overlap=150)
        final_chunks: list[tuple[str, int]] = []

        for content, page_num in pages_content:
            if not content.strip():
                continue

            # Create LlamaIndex Document abstraction
            llama_doc = LlamaDocument(
                text=content,
                metadata={
                    "doc_name": filename,
                    "page": page_num,
                    "doc_type": doc_annotation.get("document_type", "document"),
                },
            )

            # Get LlamaIndex Nodes
            nodes = splitter.get_nodes_from_documents([llama_doc])
            for node in nodes:
                node_text = node.get_content()
                if node_text.strip():
                    final_chunks.append((node_text.strip(), page_num))

        # Log formatted chunks preview
        chunk_log_dicts = [
            {
                "content": c[0],
                "page_number": c[1],
                "metadata": {
                    "doc_name": filename,
                    "page": c[1],
                    "doc_type": doc_annotation.get("document_type", "document"),
                    "short_description": doc_annotation.get("short_description", ""),
                    "summary": doc_annotation.get("summary", ""),
                },
            }
            for c in final_chunks
        ]
        IngestionLogger.log_chunks(chunk_log_dicts)

        return final_chunks, doc_annotation
